from autocorrect import Speller
from tkinter import *
from docx import *
import json
import pathlib
import threading

#all the languages as showname and speller name
allLanguages = [["en", "English"], ["pl", "Polish"], ["ru", "Russian"],
                ["uk", "Ukrainian"], ["tr", "Turkish"], ["es", "spanish"],
                ["pt", "Portuguese"], ["cs", "Czech"], ["el", "Greek"],
                ["it", "Italian"], ["fr", "French"], ["vi", "Vietnamese"],
                ["None", "None"]]

dayLanguages = {"en": "Day : ", "pl": "dzień :", "ru": "день :",
                "uk": "день :", "tr": "gün :", "es": "día :",
                "pt": "dia :", "cs": "den :", "el": "μέρα :",
                "it": "giorno :", "fr": "jour :", "vi": "ngày :",
                "None": "Day :"}

spell = Speller(lang='en')

sure = False

selectedNumber = True
nowNumb = 0

nowLanguage = []

#the Threading is needed, so it doesnt kill the main window
#*crashes anyways*
"""
def ThreadLang():
    languageChanger = threading.Thread(target = switchLang())
    languageChanger.start()
"""

#Tkinter window for selecting a language to use as autocorrector
#also error handled
def changeto(selected):
    global LanguageChanger
    global selectedLanguage
    global spell
    global nowLanguage
    nowLanguage = selected
    if selected[1] == "None":
        selectedLanguage = False
    else:
        #tries to get to the language
        try:
            spell = Speller(lang=selected[0])
        except:
            SideCMD = threading.Thread(target=putCMD)
            SideCMD.start()
            messedUP = Tk()
            MessUpInfo = Label(messedUP, text="GG you stressed it\nNow open the (Autocorrect/Location:) on your explorer and\n" +
                               "delete the messed up language in the \"data\".\nDelete the " + selected[0] + ".tar.gz")
            MessUpInfo.grid()
            messedUP.attributes("-topmost", True)
            messedUP.title("Messed up")
            messedUP.mainloop()
            
            

        selectedLanguage = True
        
    LanguageChanger.configure(text=selected[1])

#error handling for changeto() function
def putCMD():
    import os
    os.system('cmd /k "pip show autocorrect"')

#interface for changing the language
def switchLang():
 
    languageSelector = Tk()

    placements = 0
    for each in allLanguages:
        Lang = Button(languageSelector, text=each[1], command=lambda each = each:changeto(each),
                      height=3, width = 10)
        Lang.grid(row=placements//4, column=placements - (placements//4 * 4))
        placements += 1
    
    languageSelector.mainloop()

#deletes everything in the journal, usually used for developping
#purposes but can also be well used for personnal uses
def deleteall():
    global sure
    global DeleteStuff
    if sure == False:
        DeleteStuff.configure(text="Are you sure?")
    else:
        mydoc = Document()
        mydoc.save("diary.docx")
        DeleteStuff.configure(text="Deleted")
        change = open("day today.json", "w")
        change.write("")
        change.close()
        
    sure = not sure

#normal writting in a file and Json handling adding one up every day in the json.
def writeContent(content):
    global window
    
    #########Json################
    theFile = pathlib.Path("day today.json")

    if not theFile.exists():
        theFile.write_text("{}")
        dayToday = {"today": 0}

    with open('day today.json', "r") as File:
        try:
            dayToday = json.load(File)
        except:
            theFile.write_text("{}")
            dayToday = {"today": 0}
            

    with open('day today.json', 'w') as File:
        dayToday["today"] += 1
        json.dump(dayToday, File, indent=2)
    #########Json################

    #########Docx Writting################
    correctedsentance = ""
    if selectedLanguage == True:
        for eachWord in content.split(" "):
            correctedsentance += spell(eachWord) + " "
    else:
        correctedsentance = content
    info = Label(window, text="Your Update has been submitted")
    info.grid(row = 2, column = 0)
    mydoc = Document("diary.docx")
    mydoc.add_heading(dayLanguages[nowLanguage[0]] + str(dayToday["today"]), 0)
    mydoc.add_paragraph(correctedsentance)
    mydoc.save("diary.docx")
    #########Docx Writting################

#just to create a normal interface
#with all functions needed
def createWindow():
    global window
    global DeleteStuff
    global LanguageChanger
    window = Tk()
    DailyContent = Text(window, height = 30, width = 90)
    DailyContent.grid(row = 0, column = 0, columnspan = 3)
    Confirmdaily = Button(window, text="Confirm", width = 30, height = 5, command= lambda DailyContent = DailyContent:writeContent(DailyContent.get('1.0','end')), bg="Green")
    Confirmdaily.grid(row = 1, column = 0)
    LanguageChanger = Button(window, text="English", command=switchLang, width = 30, height = 5)
    LanguageChanger.grid(row = 1, column = 1)
    DeleteStuff = Button(window, text="Delete", width = 30, height = 5, command=deleteall, bg="red")
    DeleteStuff.grid(row = 1, column = 2)
    window.title("Daily journal")
    window.mainloop()

createWindow()
